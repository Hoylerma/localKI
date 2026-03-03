"""
LangChain-based RAG Pipeline Module

- Document parsing (PDF, DOCX, TXT, etc.)
- Text splitting with RecursiveCharacterTextSplitter
- Embeddings via LangChain OllamaEmbeddings
- Vector store: PGVector (Postgres + pgvector)
- rag_search Tool callable by an agent or used deterministically
"""

import asyncio
import io
import os
from typing import List

import asyncpg
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_core.tools import tool
from langchain_ollama import OllamaEmbeddings
from langchain_postgres import PGVector
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

# ---------------------------------------------------------------------------
# Configuration (env vars)
# ---------------------------------------------------------------------------

# Support both OLLAMA_BASE_URL (preferred) and legacy OLLAMA_API
OLLAMA_BASE_URL: str = os.getenv("OLLAMA_BASE_URL") or os.getenv(
    "OLLAMA_API", "http://localhost:11434"
)
EMBEDDING_MODEL: str = os.getenv("EMBEDDING_MODEL", "nomic-embed-text")
CHAT_MODEL: str = os.getenv("CHAT_MODEL", "mistral")
DATABASE_URL: str = os.getenv(
    "DATABASE_URL", "postgresql://raguser:ragpass@postgres:5432/ragdb"
)

CHUNK_SIZE: int = int(os.getenv("CHUNK_SIZE", "1000"))
CHUNK_OVERLAP: int = int(os.getenv("CHUNK_OVERLAP", "200"))
MAX_CONTEXT_CHARS: int = int(os.getenv("MAX_CONTEXT_CHARS", "4000"))
RAG_TOP_K: int = int(os.getenv("RAG_TOP_K", "5"))
MIN_SIMILARITY: float = float(os.getenv("MIN_SIMILARITY", "0.3"))

COLLECTION_NAME = "documents"
_CONTEXT_SEPARATOR = "\n\n---\n\n"

# ---------------------------------------------------------------------------
# Internal state
# ---------------------------------------------------------------------------

_vector_store: PGVector | None = None
_pool: asyncpg.Pool | None = None


def _psycopg_url() -> str:
    """Convert a standard postgresql:// URL to psycopg3 format."""
    return DATABASE_URL.replace("postgresql://", "postgresql+psycopg://", 1)


def get_vector_store() -> PGVector:
    global _vector_store
    if _vector_store is None:
        embeddings = OllamaEmbeddings(
            model=EMBEDDING_MODEL,
            base_url=OLLAMA_BASE_URL,
        )
        _vector_store = PGVector(
            embeddings=embeddings,
            collection_name=COLLECTION_NAME,
            connection=_psycopg_url(),
            use_jsonb=True,
        )
    return _vector_store


async def get_pool() -> asyncpg.Pool:
    global _pool
    if _pool is None:
        _pool = await asyncpg.create_pool(DATABASE_URL, min_size=2, max_size=10)
    return _pool


# ---------------------------------------------------------------------------
# Startup / Shutdown
# ---------------------------------------------------------------------------


async def init_db() -> None:
    """Create the pgvector extension and initialize the PGVector tables."""
    pool = await get_pool()
    async with pool.acquire() as conn:
        await conn.execute("CREATE EXTENSION IF NOT EXISTS vector;")
    # PGVector creates its own tables (langchain_pg_collection,
    # langchain_pg_embedding) on first use; trigger that now.
    loop = asyncio.get_event_loop()
    await loop.run_in_executor(None, get_vector_store)


async def close_db() -> None:
    """Release shared resources."""
    global _pool, _vector_store
    if _pool:
        await _pool.close()
        _pool = None
    _vector_store = None


# ---------------------------------------------------------------------------
# Document Parsing
# ---------------------------------------------------------------------------


def _parse_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


def _parse_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _parse_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="replace")


def parse_document(filename: str, file_bytes: bytes) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        return _parse_pdf(file_bytes)
    elif ext == "docx":
        return _parse_docx(file_bytes)
    elif ext in ("txt", "md", "csv", "json", "xml", "html"):
        return _parse_txt(file_bytes)
    else:
        raise ValueError(f"Nicht unterstütztes Dateiformat: .{ext}")


# ---------------------------------------------------------------------------
# Ingestion
# ---------------------------------------------------------------------------


async def ingest_document(filename: str, file_bytes: bytes) -> dict:
    """Parse, split, embed, and store a document via LangChain PGVector."""
    text = parse_document(filename, file_bytes)
    if not text.strip():
        raise ValueError("Dokument ist leer oder konnte nicht gelesen werden.")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )
    chunks = splitter.split_text(text)
    if not chunks:
        raise ValueError("Konnte keine Text-Chunks erstellen.")

    documents = [
        Document(
            page_content=chunk,
            metadata={"source": filename, "filename": filename, "chunk_index": i},
        )
        for i, chunk in enumerate(chunks)
    ]

    await delete_document(filename)
    vs = get_vector_store()
    # sync add_documents statt async aadd_documents
    loop = asyncio.get_event_loop()
    await loop.run_in_executor(None, vs.add_documents, documents)

    return {
        "filename": filename,
        "chunks": len(chunks),
        "characters": len(text),
    }

    # Remove previous version of this document, then add fresh chunks.
    await delete_document(filename)
    vs = get_vector_store()
    await vs.aadd_documents(documents)

    return {
        "filename": filename,
        "chunks": len(chunks),
        "characters": len(text),
    }


# ---------------------------------------------------------------------------
# Retrieval
# ---------------------------------------------------------------------------


async def rag_search_async(query: str, top_k: int = RAG_TOP_K) -> str:
    """
    Search the vector store and return a formatted context string with sources.
    Returns an empty string if nothing relevant is found.
    """
    vs = get_vector_store()
    # sync similarity_search statt async asimilarity_search
    loop = asyncio.get_event_loop()
    results = await loop.run_in_executor(
        None, lambda: vs.similarity_search_with_relevance_scores(query, k=top_k)
    )

    context_parts: List[str] = []
    total_len = 0
    for doc, score in results:
        if score < MIN_SIMILARITY:
            continue
        filename = doc.metadata.get("filename") or doc.metadata.get(
            "source", "Unbekannt"
        )
        part = f"[Quelle: {filename} | Relevanz: {score:.2f}]\n{doc.page_content}"
        if total_len + len(part) > MAX_CONTEXT_CHARS:
            remaining = MAX_CONTEXT_CHARS - total_len
            if remaining > 100:
                context_parts.append(part[:remaining] + "...")
            break
        context_parts.append(part)
        total_len += len(part) + len(_CONTEXT_SEPARATOR)

    return _CONTEXT_SEPARATOR.join(context_parts)


@tool
async def rag_search(query: str) -> str:
    """Search the knowledge base for relevant context about the query.

    Returns the most relevant text chunks together with their source filenames
    and relevance scores.  Use this tool when you need information that may be
    stored in the uploaded documents.
    """
    return await rag_search_async(query)


# ---------------------------------------------------------------------------
# Document Management
# ---------------------------------------------------------------------------


async def list_documents() -> List[dict]:
    """List all documents stored in the vector store (grouped by filename)."""
    pool = await get_pool()
    async with pool.acquire() as conn:
        rows = await conn.fetch(
            """
            SELECT
                e.cmetadata->>'filename' AS filename,
                COUNT(*) AS chunks
            FROM langchain_pg_embedding e
            JOIN langchain_pg_collection c ON e.collection_id = c.uuid
            WHERE c.name = $1
              AND e.cmetadata->>'filename' IS NOT NULL
            GROUP BY e.cmetadata->>'filename'
            ORDER BY MAX(e.id) DESC
            """,
            COLLECTION_NAME,
        )
    return [{"filename": r["filename"], "chunks": r["chunks"]} for r in rows]


async def delete_document(filename: str) -> bool:
    """Delete all chunks belonging to *filename* from the vector store."""
    pool = await get_pool()
    async with pool.acquire() as conn:
        result = await conn.execute(
            """
            DELETE FROM langchain_pg_embedding
            WHERE collection_id = (
                SELECT uuid FROM langchain_pg_collection WHERE name = $1
            )
            AND cmetadata->>'filename' = $2
            """,
            COLLECTION_NAME,
            filename,
        )
    return result != "DELETE 0"
